import io
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from app.agents.harness import harness
from app.agents.report import tools
from app.middleware.auth import get_current_user

router = APIRouter()


class PolishRequest(BaseModel):
    feedback: str


class ExportRequest(BaseModel):
    format: str = "txt"  # txt / pdf / docx


@router.post("/generate")
async def generate_report(user: dict = Depends(get_current_user)):
    result = await harness.run(
        "report",
        {"user_id": user["user_id"], "action": "generate"},
        user_id=user["user_id"],
    )
    data = result.get("data", {})
    return {"success": True, "report_text": data.get("report_text", "")}


@router.post("/polish")
async def polish_report(req: PolishRequest, user: dict = Depends(get_current_user)):
    # Load current report text
    report_text = await tools.load_report(user["user_id"])
    if not report_text:
        raise HTTPException(400, "没有已生成的报告，请先生成报告")

    result = await harness.run(
        "report",
        {
            "user_id": user["user_id"],
            "action": "polish",
            "polish_feedback": req.feedback,
            "report_text": report_text,
        },
        user_id=user["user_id"],
    )
    data = result.get("data", {})
    return {"success": True, "report_text": data.get("report_text", "")}


@router.get("/load")
async def load_report(user: dict = Depends(get_current_user)):
    report_text = await tools.load_report(user["user_id"])
    return {"success": True, "report_text": report_text or ""}


@router.post("/export")
async def export_report(req: ExportRequest, user: dict = Depends(get_current_user)):
    try:
        report_text = await tools.load_report(user["user_id"])
        if not report_text:
            raise HTTPException(400, "没有已生成的报告，请先生成报告")

        fmt = req.format.lower()
        filename = "career_report"

        if fmt == "txt":
            return StreamingResponse(
                io.BytesIO(report_text.encode("utf-8")),
                media_type="text/plain; charset=utf-8",
                headers={"Content-Disposition": f'attachment; filename="{filename}.txt"'},
            )

        if fmt == "docx":
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            style = doc.styles["Normal"]
            style.font.size = Pt(11)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing = 1.5

            for line in report_text.split("\n"):
                line = line.strip()
                if not line:
                    doc.add_paragraph("")
                    continue
                if line.startswith("【") and line.endswith("】"):
                    heading = doc.add_heading(line, level=2)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    doc.add_paragraph(line)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return StreamingResponse(
                buf,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'},
            )

        if fmt == "pdf":
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import os

            font_name = "Helvetica"
            for font_path in [
                "C:/Windows/Fonts/simsun.ttc",
                "C:/Windows/Fonts/msyh.ttc",
                "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
                "/System/Library/Fonts/STHeiti Medium.ttc",
                "/System/Library/Fonts/PingFang.ttc",
                "/Library/Fonts/Arial Unicode.ttf",
            ]:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont("ChineseFont", font_path))
                        font_name = "ChineseFont"
                        break
                    except Exception:
                        continue

            buf = io.BytesIO()
            doc_pdf = SimpleDocTemplate(buf, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm,
                                    leftMargin=2.5 * cm, rightMargin=2.5 * cm)
            styles = getSampleStyleSheet()
            normal = ParagraphStyle("NormalCN", parent=styles["Normal"], fontName=font_name,
                                    fontSize=11, leading=18, spaceAfter=6)
            heading_style = ParagraphStyle("HeadingCN", parent=styles["Heading2"], fontName=font_name,
                                     fontSize=14, leading=22, spaceAfter=10, spaceBefore=12)

            story = []
            for line in report_text.split("\n"):
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 0.3 * cm))
                    continue
                safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                if line.startswith("【") and line.endswith("】"):
                    story.append(Paragraph(safe, heading_style))
                else:
                    story.append(Paragraph(safe, normal))

            doc_pdf.build(story)
            buf.seek(0)
            return StreamingResponse(
                buf,
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{filename}.pdf"'},
            )

        raise HTTPException(400, f"不支持的导出格式: {fmt}")
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"导出失败: {str(e)}")
